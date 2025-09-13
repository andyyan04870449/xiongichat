"""
上傳處理服務
"""
import os
import uuid
import json
import csv
import mimetypes
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from datetime import datetime
import logging

from sqlalchemy import select, update
from sqlalchemy.ext.asyncio import AsyncSession
from fastapi import UploadFile, HTTPException
from PIL import Image
from PIL.ExifTags import TAGS

from app.database import get_db_context
from app.models.upload import UploadRecord, UploadStatus, UploadType, AuthoritativeMedia, AuthoritativeContacts
from app.models.knowledge import KnowledgeDocument
from app.services.embeddings import EmbeddingService
from app.services.chunker import DocumentChunker
from app.services.knowledge_manager import KnowledgeManager
from app.services.llm_service import LLMService, ContactInfo
from app.schemas.upload import (
    MediaUploadRequest, ContactsUploadRequest, ArticleUploadRequest, TextUploadRequest
)

logger = logging.getLogger(__name__)


class UploadService:
    """上傳處理服務"""
    
    def __init__(self):
        self.upload_dir = Path("uploads")
        self.upload_dir.mkdir(exist_ok=True)
        self.embedding_service = EmbeddingService()
        self.chunker = DocumentChunker()
        self.knowledge_manager = KnowledgeManager()
        self.llm_service = LLMService()  # 初始化 LLM 服務
        logger.info("UploadService initialized")
    
    async def create_upload_record(
        self, 
        filename: str, 
        upload_type: UploadType,
        file_size: Optional[int] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> UploadRecord:
        """建立上傳記錄"""
        try:
            async with get_db_context() as db:
                upload_record = UploadRecord(
                    filename=filename,
                    upload_type=upload_type,
                    status=UploadStatus.PENDING,
                    file_size=file_size,
                    metadata=metadata
                )
                db.add(upload_record)
                await db.commit()
                await db.refresh(upload_record)
                logger.info(f"Created upload record: {upload_record.id}")
                return upload_record
        except Exception as e:
            logger.error(f"Error creating upload record: {e}")
            raise HTTPException(status_code=500, detail="建立上傳記錄失敗")
    
    async def update_upload_status(
        self, 
        upload_id: uuid.UUID, 
        status: UploadStatus,
        error_message: Optional[str] = None,
        processing_log: Optional[Dict[str, Any]] = None,
        progress: Optional[int] = None
    ):
        """更新上傳狀態"""
        try:
            async with get_db_context() as db:
                update_values = {
                    "status": status,
                    "error_message": error_message,
                    "processing_log": processing_log,
                    "updated_at": datetime.utcnow()
                }
                
                # 如果提供了進度，也更新進度
                if progress is not None:
                    update_values["processing_progress"] = progress
                
                await db.execute(
                    update(UploadRecord)
                    .where(UploadRecord.id == upload_id)
                    .values(**update_values)
                )
                
                # 如果狀態是完成或失敗，更新批次任務的完成數
                if status in [UploadStatus.COMPLETED, UploadStatus.FAILED]:
                    await self._update_batch_progress(db, upload_id, status)
                
                await db.commit()
                logger.info(f"Updated upload record {upload_id} status to {status}")
        except Exception as e:
            logger.error(f"Error updating upload status: {e}")
            raise HTTPException(status_code=500, detail="更新上傳狀態失敗")
    
    async def _update_batch_progress(self, db: AsyncSession, upload_id: uuid.UUID, status: UploadStatus):
        """更新批次任務進度"""
        try:
            # 取得上傳記錄的批次ID
            result = await db.execute(
                select(UploadRecord.batch_id).where(UploadRecord.id == upload_id)
            )
            batch_id = result.scalar_one_or_none()
            
            if batch_id:
                # 取得批次任務
                from app.models.batch import BatchTask
                batch = await db.get(BatchTask, batch_id)
                if batch:
                    # 更新完成數
                    if status == UploadStatus.COMPLETED:
                        batch.completed_files = (batch.completed_files or 0) + 1
                        batch.update_status()
                        logger.info(f"Updated batch {batch_id} progress: {batch.completed_files}/{batch.total_files}")
        except Exception as e:
            logger.error(f"Error updating batch progress: {e}")
    
    async def save_uploaded_file(self, file: UploadFile, upload_id: uuid.UUID) -> str:
        """儲存上傳檔案"""
        try:
            # 建立檔案路徑
            file_extension = Path(file.filename).suffix
            file_path = self.upload_dir / f"{upload_id}{file_extension}"
            
            # 重置檔案指標到開頭
            await file.seek(0)
            
            # 儲存檔案
            with open(file_path, "wb") as buffer:
                content = await file.read()
                buffer.write(content)
            
            logger.info(f"Saved file: {file_path}, size: {len(content)} bytes")
            return str(file_path)
        except Exception as e:
            logger.error(f"Error saving file: {e}")
            raise HTTPException(status_code=500, detail="儲存檔案失敗")
    
    async def process_media_upload_with_content(
        self,
        file_content: bytes,
        filename: str,
        upload_id: uuid.UUID,
        request: MediaUploadRequest
    ) -> AuthoritativeMedia:
        """處理媒體上傳（使用已讀取的檔案內容）"""
        try:
            # 更新狀態為處理中
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING)
            
            # 儲存檔案
            file_extension = Path(filename).suffix
            file_path = self.upload_dir / f"{upload_id}{file_extension}"
            with open(file_path, "wb") as f:
                f.write(file_content)
            
            logger.info(f"Saved file: {file_path}, size: {len(file_content)} bytes")
            
            # 取得檔案資訊
            file_size = len(file_content)
            mime_type, _ = mimetypes.guess_type(filename)
            
            # 處理 EXIF 資料 (如果是圖片)
            exif_data = None
            if mime_type and mime_type.startswith('image/'):
                try:
                    with Image.open(file_path) as img:
                        exif_dict = img._getexif()
                        if exif_dict:
                            exif_data = {}
                            for tag_id, value in exif_dict.items():
                                tag = TAGS.get(tag_id, tag_id)
                                exif_data[tag] = str(value)
                except Exception as e:
                    logger.warning(f"Could not extract EXIF data: {e}")
            
            # 儲存到資料庫
            async with get_db_context() as db:
                media_record = AuthoritativeMedia(
                    upload_id=upload_id,
                    filename=filename,
                    file_path=str(file_path),
                    file_size=file_size,
                    mime_type=mime_type,
                    tags=request.get("tags", []),
                    description=request.get("description"),
                    exif_data=exif_data
                )
                db.add(media_record)
                await db.commit()
                await db.refresh(media_record)
            
            # 建立 RAG 索引
            await self._create_media_rag_index(media_record)
            
            # 更新狀態為完成
            await self.update_upload_status(
                upload_id, 
                UploadStatus.COMPLETED,
                processing_log={"media_id": str(media_record.id)}
            )
            
            logger.info(f"Processed media upload: {media_record.id}")
            return media_record
            
        except Exception as e:
            logger.error(f"Error processing media upload: {e}")
            await self.update_upload_status(upload_id, UploadStatus.FAILED, str(e))
            raise HTTPException(status_code=500, detail=f"處理媒體上傳失敗: {str(e)}")

    async def process_media_upload(
        self, 
        file: UploadFile, 
        upload_id: uuid.UUID,
        request: MediaUploadRequest
    ) -> AuthoritativeMedia:
        """處理媒體上傳"""
        try:
            # 更新狀態為處理中
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING)
            
            # 儲存檔案
            file_path = await self.save_uploaded_file(file, upload_id)
            
            # 取得檔案資訊
            file_size = os.path.getsize(file_path)
            mime_type, _ = mimetypes.guess_type(file.filename)
            
            # 處理 EXIF 資料 (如果是圖片)
            exif_data = None
            if mime_type and mime_type.startswith('image/'):
                try:
                    with Image.open(file_path) as img:
                        exif_dict = img._getexif()
                        if exif_dict:
                            exif_data = {}
                            for tag_id, value in exif_dict.items():
                                tag = TAGS.get(tag_id, tag_id)
                                exif_data[tag] = str(value)
                except Exception as e:
                    logger.warning(f"Could not extract EXIF data: {e}")
            
            # 儲存到資料庫
            async with get_db_context() as db:
                media_record = AuthoritativeMedia(
                    upload_id=upload_id,
                    filename=file.filename,
                    file_path=file_path,
                    file_size=file_size,
                    mime_type=mime_type,
                    tags=request.tags,
                    description=request.description,
                    exif_data=exif_data
                )
                db.add(media_record)
                await db.commit()
                await db.refresh(media_record)
            
            # 建立 RAG 索引
            await self._create_media_rag_index(media_record)
            
            # 更新狀態為完成
            await self.update_upload_status(
                upload_id, 
                UploadStatus.COMPLETED,
                processing_log={"media_id": str(media_record.id)}
            )
            
            logger.info(f"Processed media upload: {media_record.id}")
            return media_record
            
        except Exception as e:
            logger.error(f"Error processing media upload: {e}")
            await self.update_upload_status(upload_id, UploadStatus.FAILED, str(e))
            raise HTTPException(status_code=500, detail=f"處理媒體上傳失敗: {str(e)}")
    
    async def process_contacts_upload_with_content(
        self,
        file_content: bytes,
        filename: str,
        upload_id: uuid.UUID,
        request: ContactsUploadRequest
    ) -> List[AuthoritativeContacts]:
        """處理聯絡人上傳（使用已讀取的檔案內容）"""
        try:
            # 更新狀態為處理中，進度 10%
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING, progress=10)
            
            # 儲存檔案，進度 20%
            file_extension = Path(filename).suffix.lower()
            file_path = self.upload_dir / f"{upload_id}{file_extension}"
            with open(file_path, "wb") as f:
                f.write(file_content)
            
            logger.info(f"Saved file: {file_path}, size: {len(file_content)} bytes")
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING, progress=20)
            
            # 提取文字內容，進度 30%
            text_content = await self._extract_text_from_file(str(file_path))
            logger.info(f"Extracted text length: {len(text_content)} characters")
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING, progress=30)
            
            # 使用 LLM 辨識聯絡人資訊，進度 50%
            logger.info("Using LLM to extract contacts from text...")
            contact_infos = await self.llm_service.extract_contacts_from_text(text_content)
            logger.info(f"LLM extracted {len(contact_infos)} contacts")
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING, progress=50)
            
            # 儲存到資料庫，進度 70%
            contacts_records = []
            async with get_db_context() as db:
                for contact_info in contact_infos:
                    # 資料標準化
                    phone = self.llm_service.standardize_phone(contact_info.phone)
                    address = self.llm_service.standardize_address(contact_info.address)
                    
                    # 建立資料庫記錄
                    contact_record = AuthoritativeContacts(
                        upload_id=upload_id,
                        name=contact_info.name,
                        organization=contact_info.name,  # 保留相容性
                        category=contact_info.category,
                        phone=phone,
                        email=contact_info.email,
                        address=address,
                        services=contact_info.services,
                        contact_person=contact_info.contact_person,
                        tags=[contact_info.category] if contact_info.category else [],
                        notes=contact_info.notes
                    )
                    db.add(contact_record)
                    contacts_records.append(contact_record)
                
                await db.commit()
                
                for record in contacts_records:
                    await db.refresh(record)
            
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING, progress=70)
            
            # 建立 RAG 索引和向量，進度 90%
            for contact_record in contacts_records:
                # 生成增強描述
                description = await self.llm_service.enhance_contact_description(
                    ContactInfo(
                        name=contact_record.name,
                        category=contact_record.category,
                        phone=contact_record.phone,
                        address=contact_record.address,
                        services=contact_record.services,
                        email=contact_record.email,
                        contact_person=contact_record.contact_person,
                        notes=contact_record.notes
                    )
                )
                
                # 建立向量索引
                await self._create_contact_rag_index_with_description(contact_record, description)
            
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING, progress=90)
            
            # 更新狀態為完成，進度 100%
            await self.update_upload_status(
                upload_id, 
                UploadStatus.COMPLETED,
                processing_log={
                    'contacts_count': len(contacts_records),
                    'llm_model': self.llm_service.model
                },
                progress=100
            )
            
            logger.info(f"Processed contacts upload: {len(contacts_records)} contacts")
            return contacts_records
            
        except Exception as e:
            logger.error(f"Error processing contacts upload: {e}")
            await self.update_upload_status(upload_id, UploadStatus.FAILED, str(e))
            raise HTTPException(status_code=500, detail=f"處理聯絡人上傳失敗: {str(e)}")
    
    async def process_contacts_upload(
        self, 
        file: UploadFile, 
        upload_id: uuid.UUID,
        request: ContactsUploadRequest
    ) -> List[AuthoritativeContacts]:
        """處理聯絡人上傳"""
        try:
            # 更新狀態為處理中
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING)
            
            # 儲存檔案
            file_path = await self.save_uploaded_file(file, upload_id)
            
            # 解析 CSV/Excel 檔案
            contacts_data = await self._parse_contacts_file(file_path, request.field_mapping)
            
            # 儲存到資料庫
            contacts_records = []
            async with get_db_context() as db:
                for contact_data in contacts_data:
                    # 確保 organization 欄位存在
                    if 'organization' not in contact_data or not contact_data['organization']:
                        logger.warning(f"Skipping contact without organization: {contact_data}")
                        continue
                    
                    contact_record = AuthoritativeContacts(
                        upload_id=upload_id,
                        organization=contact_data['organization'],
                        phone=contact_data.get('phone'),
                        email=contact_data.get('email'),
                        address=contact_data.get('address'),
                        tags=contact_data.get('tags'),
                        notes=contact_data.get('notes')
                    )
                    db.add(contact_record)
                    contacts_records.append(contact_record)
                
                await db.commit()
                
                for record in contacts_records:
                    await db.refresh(record)
            
            # 建立 RAG 索引
            for contact_record in contacts_records:
                await self._create_contact_rag_index(contact_record)
            
            # 更新狀態為完成
            await self.update_upload_status(
                upload_id, 
                UploadStatus.COMPLETED,
                processing_log={"contacts_count": len(contacts_records)}
            )
            
            logger.info(f"Processed contacts upload: {len(contacts_records)} contacts")
            return contacts_records
            
        except Exception as e:
            logger.error(f"Error processing contacts upload: {e}")
            await self.update_upload_status(upload_id, UploadStatus.FAILED, str(e))
            raise HTTPException(status_code=500, detail=f"處理聯絡人上傳失敗: {str(e)}")
    
    async def process_article_upload_with_content(
        self,
        file_content: bytes,
        filename: str,
        upload_id: uuid.UUID,
        request: ArticleUploadRequest
    ) -> KnowledgeDocument:
        """處理文章上傳（使用已讀取的檔案內容）"""
        try:
            # 更新狀態為處理中，進度 10%
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING, progress=10)
            
            # 儲存檔案，進度 30%
            file_extension = Path(filename).suffix
            file_path = self.upload_dir / f"{upload_id}{file_extension}"
            with open(file_path, "wb") as f:
                f.write(file_content)
            
            logger.info(f"Saved file: {file_path}, size: {len(file_content)} bytes")
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING, progress=30)
            
            # 解析文件內容，進度 50%
            content = await self._extract_text_from_file(str(file_path))
            title = Path(filename).stem  # 使用檔名作為標題
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING, progress=50)
            
            # 建立知識文件，進度 80%
            document_id = await self.knowledge_manager.add_document(
                title=title,
                content=content,
                source=request.source,
                category=request.category,
                lang=request.lang,
                published_date=request.published_date
            )
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING, progress=80)
            
            # 更新狀態為完成，進度 100%
            await self.update_upload_status(
                upload_id, 
                UploadStatus.COMPLETED,
                processing_log={"document_id": document_id},
                progress=100
            )
            
            logger.info(f"Processed article upload: {document_id}")
            
            # 建立並返回文件物件
            from app.models.knowledge import KnowledgeDocument
            document = KnowledgeDocument(id=document_id, title=title, content=content)
            return document
            
        except Exception as e:
            logger.error(f"Error processing article upload: {e}")
            await self.update_upload_status(upload_id, UploadStatus.FAILED, str(e))
            raise HTTPException(status_code=500, detail=f"處理文章上傳失敗: {str(e)}")

    async def process_article_upload(
        self, 
        file: UploadFile, 
        upload_id: uuid.UUID,
        request: ArticleUploadRequest
    ) -> KnowledgeDocument:
        """處理文章上傳"""
        try:
            # 更新狀態為處理中，進度 10%
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING, progress=10)
            
            # 儲存檔案，進度 30%
            file_path = await self.save_uploaded_file(file, upload_id)
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING, progress=30)
            
            # 解析文件內容，進度 50%
            content = await self._extract_text_from_file(file_path)
            title = Path(file.filename).stem  # 使用檔名作為標題
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING, progress=50)
            
            # 建立知識文件，進度 80%
            document_id = await self.knowledge_manager.add_document(
                title=title,
                content=content,
                source=request.source,
                category=request.category,
                lang=request.lang,
                published_date=request.published_date
            )
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING, progress=80)
            
            # 更新狀態為完成，進度 100%
            await self.update_upload_status(
                upload_id, 
                UploadStatus.COMPLETED,
                processing_log={"document_id": document_id},
                progress=100
            )
            
            logger.info(f"Processed article upload: {document_id}")
            
            # 建立並返回文件物件
            from app.models.knowledge import KnowledgeDocument
            document = KnowledgeDocument(id=document_id, title=title, content=content)
            return document
            
        except Exception as e:
            logger.error(f"Error processing article upload: {e}")
            await self.update_upload_status(upload_id, UploadStatus.FAILED, str(e))
            raise HTTPException(status_code=500, detail=f"處理文章上傳失敗: {str(e)}")
    
    async def process_text_upload(
        self, 
        upload_id: uuid.UUID,
        request: TextUploadRequest
    ) -> KnowledgeDocument:
        """處理文字上傳"""
        try:
            # 更新狀態為處理中
            await self.update_upload_status(upload_id, UploadStatus.PROCESSING)
            
            # 建立知識文件
            document_id = await self.knowledge_manager.add_document(
                title=request.title,
                content=request.content,
                source=request.source,
                category=request.category,
                lang=request.lang,
                published_date=request.published_date
            )
            
            # 更新狀態為完成
            await self.update_upload_status(
                upload_id, 
                UploadStatus.COMPLETED,
                processing_log={"document_id": document_id}
            )
            
            logger.info(f"Processed text upload: {document_id}")
            
            # 建立並返回文件物件
            from app.models.knowledge import KnowledgeDocument
            document = KnowledgeDocument(id=document_id, title=request.title, content=request.content)
            return document
            
        except Exception as e:
            logger.error(f"Error processing text upload: {e}")
            await self.update_upload_status(upload_id, UploadStatus.FAILED, str(e))
            raise HTTPException(status_code=500, detail=f"處理文字上傳失敗: {str(e)}")
    
    async def _create_media_rag_index(self, media_record: AuthoritativeMedia):
        """為媒體建立 RAG 索引"""
        try:
            # 建立關鍵字摘要
            keywords = []
            if media_record.tags:
                keywords.extend(media_record.tags)
            if media_record.description:
                keywords.append(media_record.description)
            
            # 建立摘要文字
            summary_text = f"媒體檔案: {media_record.filename}"
            if keywords:
                summary_text += f" 標籤: {', '.join(keywords)}"
            if media_record.description:
                summary_text += f" 描述: {media_record.description}"
            
            # 建立知識文件 (用於 RAG 索引)
            document = await self.knowledge_manager.add_document(
                title=f"媒體: {media_record.filename}",
                content=summary_text,
                source="authoritative_media",
                category="media",
                lang="zh-TW"
            )
            
            logger.info(f"Created RAG index for media: {media_record.id}")
            
        except Exception as e:
            logger.error(f"Error creating media RAG index: {e}")
    
    async def _create_contact_rag_index(self, contact_record: AuthoritativeContacts):
        """為機構聯絡資訊建立 RAG 索引"""
        try:
            # 建立關鍵字摘要
            keywords = [contact_record.organization]
            if contact_record.tags:
                keywords.extend(contact_record.tags)
            
            # 建立摘要文字
            summary_text = f"機構聯絡資訊: {contact_record.organization}"
            
            # 添加聯絡方式說明（不直接包含實際號碼）
            contact_methods = []
            if contact_record.phone:
                contact_methods.append("提供電話聯絡")
            if contact_record.email:
                contact_methods.append("提供電子郵件")
            if contact_record.address:
                contact_methods.append("提供地址資訊")
            
            if contact_methods:
                summary_text += f" 聯絡方式: {'、'.join(contact_methods)}"
            
            if contact_record.notes:
                summary_text += f" 服務說明: {contact_record.notes}"
            
            if keywords:
                summary_text += f" 關鍵字: {', '.join(keywords)}"
            
            # 建立知識文件 (用於 RAG 索引)
            document = await self.knowledge_manager.add_document(
                title=f"機構聯絡資訊: {contact_record.organization}",
                content=summary_text,
                source="authoritative_contacts",
                category="contacts",
                lang="zh-TW"
            )
            
            logger.info(f"Created RAG index for contact: {contact_record.id}")
            
        except Exception as e:
            logger.error(f"Error creating contact RAG index: {e}")
    
    async def _create_contact_rag_index_with_description(
        self, 
        contact_record: AuthoritativeContacts,
        description: str
    ):
        """為機構聯絡資訊建立增強的 RAG 索引"""
        try:
            # 建立知識文件 (用於 RAG 索引)
            document_id = await self.knowledge_manager.add_document(
                title=f"機構: {contact_record.name or contact_record.organization}",
                content=description,
                source="authoritative_contacts",
                category=contact_record.category or "contacts",
                lang="zh-TW"
            )
            
            logger.info(f"Created enhanced RAG index for contact: {contact_record.id}, document: {document_id}")
            
        except Exception as e:
            logger.error(f"Error creating enhanced contact RAG index: {e}")
    
    async def _parse_contacts_file(
        self, 
        file_path: str, 
        field_mapping: Dict[str, str]
    ) -> List[Dict[str, Any]]:
        """解析聯絡人檔案"""
        contacts_data = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                # 嘗試自動偵測分隔符號
                sample = file.read(1024)
                file.seek(0)
                
                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter
                
                reader = csv.DictReader(file, delimiter=delimiter)
                
                for row in reader:
                    contact_data = {}
                    
                    # 根據欄位映射轉換資料
                    for db_field, csv_field in field_mapping.items():
                        if csv_field in row and row[csv_field]:
                            value = row[csv_field].strip()
                            # 特殊處理 tags 欄位，轉換成列表
                            if db_field == 'tags' and value:
                                contact_data[db_field] = [tag.strip() for tag in value.split(',')]
                            else:
                                contact_data[db_field] = value
                    
                    # 驗證必要欄位：organization 是必填，phone 或 email 至少要有一個
                    if contact_data.get('organization') and \
                       (contact_data.get('phone') or contact_data.get('email')):
                        contacts_data.append(contact_data)
                
                logger.info(f"Parsed {len(contacts_data)} contacts from file")
                return contacts_data
                
        except Exception as e:
            logger.error(f"Error parsing contacts file: {e}")
            raise HTTPException(status_code=400, detail=f"解析聯絡人檔案失敗: {str(e)}")
    
    async def _extract_text_from_file(self, file_path: str) -> str:
        """從檔案中提取文字"""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
                    
            elif file_extension == '.csv':
                # 讀取 CSV 內容
                text_content = []
                with open(file_path, 'r', encoding='utf-8-sig') as file:
                    csv_reader = csv.reader(file)
                    for row in csv_reader:
                        text_content.append(' | '.join(row))
                return '\n'.join(text_content)
                
            elif file_extension in ['.xlsx', '.xls']:
                # 使用 openpyxl 讀取 Excel（需要安裝）
                try:
                    import openpyxl
                    workbook = openpyxl.load_workbook(file_path, read_only=True)
                    text_content = []
                    for sheet_name in workbook.sheetnames:
                        sheet = workbook[sheet_name]
                        text_content.append(f"工作表: {sheet_name}")
                        for row in sheet.iter_rows(values_only=True):
                            if any(row):  # 跳過空行
                                text_content.append(' | '.join(str(cell) if cell else '' for cell in row))
                    workbook.close()
                    return '\n'.join(text_content)
                except ImportError:
                    logger.warning("openpyxl not installed, treating Excel as CSV")
                    # 降級為簡單文字提取
                    with open(file_path, 'rb') as file:
                        content = file.read()
                        return content.decode('utf-8', errors='ignore')
                        
            elif file_extension == '.json':
                # 讀取 JSON 並格式化為文字
                with open(file_path, 'r', encoding='utf-8') as file:
                    data = json.load(file)
                    return json.dumps(data, ensure_ascii=False, indent=2)
                    
            elif file_extension == '.pdf':
                # 使用 PyPDF2 讀取 PDF（需要安裝）
                try:
                    import PyPDF2
                    text_content = []
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        for page_num in range(len(pdf_reader.pages)):
                            page = pdf_reader.pages[page_num]
                            text_content.append(page.extract_text())
                    return '\n'.join(text_content)
                except ImportError:
                    logger.warning("PyPDF2 not installed, cannot extract PDF text")
                    return f"PDF 檔案: {Path(file_path).name} (需要安裝 PyPDF2 來提取內容)"
                    
            elif file_extension in ['.doc', '.docx']:
                # 使用 python-docx 讀取 Word 文檔（需要安裝）
                try:
                    import docx
                    doc = docx.Document(file_path)
                    text_content = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(paragraph.text)
                    # 也提取表格內容
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                            if row_text.strip():
                                text_content.append(row_text)
                    return '\n'.join(text_content)
                except ImportError:
                    logger.warning("python-docx not installed, cannot extract Word text")
                    return f"Word 檔案: {Path(file_path).name} (需要安裝 python-docx 來提取內容)"
                    
            else:
                # 對於其他格式，嘗試以文字模式讀取
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        return file.read()
                except UnicodeDecodeError:
                    # 如果無法解碼，返回檔案基本資訊
                    return f"檔案: {Path(file_path).name} (無法提取文字內容)"
                
        except Exception as e:
            logger.error(f"Error extracting text from file: {e}")
            raise HTTPException(status_code=400, detail=f"提取檔案文字失敗: {str(e)}")
    
    async def get_upload_record(self, upload_id: uuid.UUID) -> Optional[UploadRecord]:
        """取得上傳記錄"""
        try:
            async with get_db_context() as db:
                result = await db.execute(
                    select(UploadRecord).where(UploadRecord.id == upload_id)
                )
                return result.scalar_one_or_none()
        except Exception as e:
            logger.error(f"Error getting upload record: {e}")
            return None
    
    async def get_recent_uploads(self, limit: int = 10) -> List[UploadRecord]:
        """取得最近上傳記錄"""
        try:
            async with get_db_context() as db:
                result = await db.execute(
                    select(UploadRecord)
                    .order_by(UploadRecord.created_at.desc())
                    .limit(limit)
                )
                return result.scalars().all()
        except Exception as e:
            logger.error(f"Error getting recent uploads: {e}")
            return []
